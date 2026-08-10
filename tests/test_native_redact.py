"""Tests for native PDF/DOCX redaction and surface helpers."""

from __future__ import annotations

from pathlib import Path

import pytest

from anonymizer.anonymize.surfaces import (
    RedactSurface,
    surface_search_variants,
    surfaces_from_mapping,
)
from anonymizer.output.docx_redact import redact_docx
from anonymizer.output.native import (
    normalize_output_format,
    wants_markdown,
    wants_native,
    write_native_redacted,
)
from anonymizer.output.pdf_redact import redact_pdf


def test_surfaces_longest_first() -> None:
    m = {
        "[ORG_1]": "ACME",
        "[ORG_2]": "ACME Corp",
        "[PERSON_1]": "Alice",
    }
    surfs = surfaces_from_mapping(m)
    clears = [s.clear for s in surfs]
    assert clears[0] == "ACME Corp"  # longest first
    assert set(clears) == {"ACME Corp", "ACME", "Alice"}


def test_surface_variants_nbsp() -> None:
    v = surface_search_variants("Foo\u00a0Bar")
    assert "Foo\u00a0Bar" in v
    assert "Foo Bar" in v


def test_normalize_output_format() -> None:
    assert normalize_output_format(None) == "md"
    assert normalize_output_format("native") == "source"
    assert normalize_output_format("BOTH") == "both"
    with pytest.raises(ValueError):
        normalize_output_format("excel")


def test_wants_flags() -> None:
    assert wants_markdown("md") and not wants_native("md")
    assert wants_native("source") and not wants_markdown("source")
    assert wants_markdown("both") and wants_native("both")


def test_pdf_redact_removes_cleartext(tmp_path: Path) -> None:
    import pymupdf as fitz

    src = tmp_path / "in.pdf"
    dest = tmp_path / "out.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Contact Alice Wonderland at HQ.")
    doc.save(src)
    doc.close()

    surfaces = [
        RedactSurface(clear="Alice Wonderland", placeholder="[PERSON_1]"),
    ]
    stats = redact_pdf(src, surfaces, dest)
    assert stats.surfaces_found == 1
    assert stats.hit_count >= 1
    assert dest.is_file()

    after = fitz.open(dest)
    text = after[0].get_text()
    after.close()
    assert "Alice Wonderland" not in text


def test_docx_redact_placeholder(tmp_path: Path) -> None:
    from docx import Document

    src = tmp_path / "in.docx"
    dest = tmp_path / "out.docx"
    d = Document()
    d.add_paragraph("Signed by Alice Wonderland on behalf of ACME Corp.")
    d.save(src)

    surfaces = [
        RedactSurface(clear="Alice Wonderland", placeholder="[PERSON_1]"),
        RedactSurface(clear="ACME Corp", placeholder="[ORG_1]"),
    ]
    stats = redact_docx(src, surfaces, dest, style="placeholder")
    assert stats.surfaces_found == 2
    assert dest.is_file()

    out = Document(str(dest))
    body = "\n".join(p.text for p in out.paragraphs)
    assert "Alice Wonderland" not in body
    assert "[PERSON_1]" in body
    assert "[ORG_1]" in body


def test_docx_redact_remove(tmp_path: Path) -> None:
    from docx import Document

    src = tmp_path / "in.docx"
    dest = tmp_path / "out.docx"
    d = Document()
    d.add_paragraph("Name: Alice Wonderland.")
    d.save(src)

    surfaces = [RedactSurface(clear="Alice Wonderland", placeholder="[PERSON_1]")]
    redact_docx(src, surfaces, dest, style="remove")
    body = "\n".join(p.text for p in Document(str(dest)).paragraphs)
    assert "Alice Wonderland" not in body
    assert "[PERSON_1]" not in body


def test_write_native_dispatch(tmp_path: Path) -> None:
    import pymupdf as fitz

    src = tmp_path / "x.pdf"
    dest = tmp_path / "x.anonymized.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Secret Bob Builder")
    doc.save(src)
    doc.close()

    stats = write_native_redacted(
        src, dest, {"[PERSON_1]": "Bob Builder"}, redact_style="placeholder"
    )
    assert stats is not None
    assert stats.format == "pdf"
    assert dest.is_file()
