"""Write a redacted DOCX by replacing cleartext in paragraphs and tables."""

from __future__ import annotations

import logging
import re
from pathlib import Path

from anonymizer.anonymize.surfaces import RedactSurface, surface_search_variants
from anonymizer.output.native_stats import NativeRedactStats

logger = logging.getLogger(__name__)


def _replace_in_text(text: str, clear: str, replacement: str) -> tuple[str, int]:
    """Replace all occurrences of *clear* (and simple variants) in *text*."""
    count = 0
    out = text
    for variant in surface_search_variants(clear):
        if not variant or variant not in out:
            continue
        n = out.count(variant)
        if n:
            out = out.replace(variant, replacement)
            count += n
    return out, count


def _set_paragraph_text(paragraph, new_text: str) -> None:
    """Replace full paragraph text, keeping first run's style when possible."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _iter_all_paragraphs(doc):
    """Body paragraphs, table cells, headers, and footers."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for part in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            if part is None:
                continue
            try:
                for p in part.paragraphs:
                    yield p
                for table in part.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                yield p
            except Exception:  # noqa: BLE001 — linked section parts
                continue


def redact_docx(
    source: Path,
    surfaces: list[RedactSurface],
    dest: Path,
    *,
    style: str = "placeholder",
) -> NativeRedactStats:
    """Copy *source* to *dest* with cleartext replaced per *style*.

    *style* ``placeholder`` inserts ``[TYPE_n]`` tags; ``remove`` deletes text.
    Cross-run matches are handled by rewriting the full paragraph text into
    the first run (formatting of later runs may be lost for that paragraph).
    """
    from docx import Document

    source = Path(source)
    dest = Path(dest)
    stats = NativeRedactStats(format="docx", surfaces_total=len(surfaces))
    if style not in ("placeholder", "remove"):
        style = "placeholder"

    doc = Document(str(source))
    if not surfaces:
        dest.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(dest))
        stats.output_path = str(dest)
        return stats

    found: set[str] = set()
    for paragraph in _iter_all_paragraphs(doc):
        full = paragraph.text
        if not full:
            continue
        new_full = full
        para_hits = 0
        for surface in surfaces:
            replacement = "" if style == "remove" else surface.placeholder
            new_full, n = _replace_in_text(new_full, surface.clear, replacement)
            if n:
                found.add(surface.clear)
                para_hits += n
                stats.hit_count += n
        if para_hits and new_full != full:
            if style == "remove":
                new_full = re.sub(r"[^\S\n]{2,}", " ", new_full)
            _set_paragraph_text(paragraph, new_full)

    for surface in surfaces:
        if surface.clear in found:
            stats.surfaces_found += 1
        else:
            stats.surfaces_missed += 1
            stats.missed.append(surface.clear)

    # Best-effort document property scrub (Author, Title, …)
    try:
        props = doc.core_properties
        for attr in (
            "author",
            "category",
            "comments",
            "content_status",
            "identifier",
            "keywords",
            "last_modified_by",
            "subject",
            "title",
        ):
            try:
                setattr(props, attr, None)
            except (AttributeError, ValueError, TypeError):
                try:
                    setattr(props, attr, "")
                except (AttributeError, ValueError, TypeError):
                    pass
    except Exception as exc:  # noqa: BLE001
        logger.debug("core_properties scrub failed: %s", exc)

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    stats.output_path = str(dest)
    return stats
